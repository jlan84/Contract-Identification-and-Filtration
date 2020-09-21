import PyPDF2
import pdftotext as ptt
from directoryassistor import DirectoryAssistor
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import os
from docx import Document
import re
import numpy as np

class ImageConverter():

    def __init__(self, directory):
        self.directory = directory
        self.ds = DirectoryAssistor()
    
    def convert_image(self, file_name, file_type='.pdf', dpi=500):
        '''
        Converts a file of type file_type to .txt using OCR

        Params
        file_name: string that is the name of the file that needs to be read in
        file_type: str that is the type of file being read in
        dpi: int that is the dots per inch of the file being read in

        '''
        pages = convert_from_path(self.directory+file_name, dpi=dpi)
        image_counter = 1
        image_names = []
               
        for page in pages:
            image_name = 'page_'+str(image_counter)+'.jpg'
            image_names.append(image_name)
            page.save(image_name, 'JPEG')
            image_counter += 1

        new_file_name = file_name.replace(file_type, '.txt')
        filelimit = image_counter - 1
        outfile = self.directory+new_file_name

        f = open(outfile, 'a')
        for i in range(1, filelimit+1):
            image_name = "page_"+str(i)+".jpg"
            text = str(((pytesseract.image_to_string(Image.open(image_name)))))
            text = text.replace('-\n', '')
            f.write(text)
        f.close()

        for img in image_names:
            self.ds.delete_files(img)
        self.ds.delete_files(directory+file_name)
    
    def convert_txt_to_doc(self, text_file):
        '''
        Converts a .txt document to a .doc format

        Params
        text_file: name of .txt file stored in the directory for the object

        '''

        document = Document()
        new_name = text_file.replace('.txt','')
        document.add_heading(new_name, 0)

        myfile = open(self.directory+text_file).read()
        myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile)
        p = document.add_paragraph(myfile)
        document.save(directory+new_name+'.doc')
    
    
    
    

if __name__ == '__main__':
    
    directory = '/Users/justinlansdale/Documents/Galvanize/Capstone2/\
Contract-Classifier/data/CityofChicago/ChicagoContracts/'

    ic = ImageConverter(directory)
    ic.convert_image('document.pdf')
    ic.convert_txt_to_doc('document.txt')
    